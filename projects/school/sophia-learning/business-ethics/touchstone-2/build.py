from docx import Document
from docx.shared import Pt, Inches
import os

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

def add_bold_label(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(12)
    run2 = p.add_run(text)
    run2.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(12)
    return p

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

# Header
p = doc.add_paragraph()
run = p.add_run("Touchstone 2: Review a Code of Ethics")
run.bold = True
run.font.size = Pt(14)

add_para(doc, "Name: Jaylen Davis")
add_para(doc, "Date: April 19, 2026")
add_para(doc, "Course: Business Ethics")

add_heading(doc, "Question 1: Company Selected")
add_para(doc, "The company whose code of ethics I have chosen is Starbucks Coffee Company. Their code, titled the Standards of Business Conduct, is publicly available and outlines the ethical expectations for all partners worldwide. Starbucks is one of the most recognized brands globally, making their code of ethics a useful and substantive document to analyze.")

add_heading(doc, "Question 2: Policy Summaries by Topic")

add_bold_label(doc, "Basic Conduct: ", "Starbucks expects all partners to act with honesty, integrity, and mutual respect in every aspect of their work. The code makes clear that partners should treat one another, customers, and business contacts with dignity regardless of background or position.")

add_bold_label(doc, "Confidentiality: ", "Partners are required to protect confidential business information and may not share proprietary data, trade secrets, or internal strategies with unauthorized individuals inside or outside the company. The code emphasizes that this obligation continues even after a partner leaves the company.")

add_bold_label(doc, "Competition: ", "Starbucks requires that all business be conducted fairly and in compliance with applicable antitrust and competition laws. Partners are prohibited from engaging in anticompetitive practices such as price fixing, market allocation, or bid rigging with competitors.")

add_bold_label(doc, "Giving and Receiving Gifts: ", "Partners may not accept gifts, meals, entertainment, or other benefits from suppliers, vendors, or others doing business with Starbucks if doing so could create a conflict of interest or appear to influence business decisions.")

add_bold_label(doc, "Insider Trading: ", "Partners who have access to material nonpublic information about Starbucks or its business partners may not trade in company stock or disclose that information to others until it has been properly released to the public.")

add_bold_label(doc, "Social Media Usage: ", "Partners are expected to be thoughtful and professional in their personal social media activity, ensuring they do not reveal confidential information, make false statements about the company, or post content that could harm the Starbucks brand or create a hostile environment for others.")

add_bold_label(doc, "Disparagement: ", "Partners are expected to speak respectfully about Starbucks, competitors, and business partners. Making false, misleading, or derogatory statements about other companies or individuals is inconsistent with the company's values and professional standards.")

add_heading(doc, "Question 3: Ethical Minimum vs. Higher Ethical Standard")
add_para(doc, "The Starbucks Standards of Business Conduct reflect a combination of ethical minimums and a higher ethical standard depending on the policy area. Certain provisions, such as the prohibition on insider trading and the requirement to comply with antitrust law, represent ethical minimums because they are legally mandated by federal law and failure to comply would result in criminal or civil liability. Similarly, the non-discrimination commitment reflects at minimum a legal obligation under the Civil Rights Act of 1964 and related employment statutes. However, Starbucks goes well beyond these legal baselines in several meaningful ways. Their commitment to treating all employees as partners and fostering a culture of belonging is a voluntary decision that exceeds what the law demands. Their ethical sourcing commitments and community investment programs represent a higher ethical standard grounded in corporate social responsibility rather than regulatory compliance. The overall tone of the code encourages partners to do the right thing even in ambiguous situations where no rule directly applies, which reflects a higher standard rooted in personal integrity. This blend of legal minimums and aspirational values illustrates how a mature ethics program operates across both floors and ceilings of ethical behavior.")

add_heading(doc, "Question 4: Organizational Culture (Competing Values Framework)")
add_para(doc, "Based on the Starbucks Standards of Business Conduct, the company demonstrates a primarily clan culture within the Competing Values Framework. The most telling indicator is the consistent use of the term partners to refer to all employees, which signals a collaborative, family-like orientation that prioritizes people and relationships over strict hierarchy or financial metrics. The code explicitly emphasizes mutual respect, open communication, and a commitment to inclusivity, all of which are core characteristics of clan culture. Partners are encouraged to raise ethical concerns through open dialogue and accessible channels rather than through rigid chains of command, reinforcing a trust-based internal environment. That said, elements of hierarchy culture are also visible, particularly in the structured compliance requirements governing financial reporting, legal obligations, and insider trading restrictions. These sections reflect a need for formal controls and standardized processes associated with hierarchy culture. However, the dominant voice of the document is relational and values-driven, and the people-centered language throughout suggests that clan culture is the primary orientation Starbucks aims to cultivate.")

add_heading(doc, "Question 5: Ethical Theory")
add_para(doc, "The Starbucks code of ethics most closely aligns with virtue ethics, with secondary elements of stakeholder theory also present throughout the document. Virtue ethics holds that ethical behavior stems from internalized character traits such as honesty, fairness, and integrity, rather than from calculating outcomes or following rigid rules. Starbucks reflects this approach by consistently calling on partners to act with integrity and to do what is right even in ambiguous situations, rather than simply enumerating a list of prohibited behaviors. For example, the emphasis on honesty in all communications and mutual respect for all people regardless of position reflects the virtue ethics ideal that good character is the foundation of ethical conduct. At the same time, the code acknowledges Starbucks' responsibility to a broad range of stakeholders including employees, customers, suppliers, communities, and shareholders, which aligns with stakeholder theory and its emphasis on balancing competing interests. This is evident in their ethical sourcing commitments and community investment language throughout the document. Together, these two frameworks suggest that Starbucks views ethics as both character-driven and relationship-centered, with virtue ethics serving as the primary theoretical foundation.")

add_heading(doc, "Question 6: Personal Reflection")

add_bold_label(doc, "a. Reasonable and Employee-Centered: ", "The Starbucks code of ethics is largely reasonable and notably employee-centered. The deliberate decision to refer to all employees as partners sends a clear signal that the company views its workforce as genuine stakeholders in the organization rather than interchangeable labor. The code's emphasis on mutual respect, inclusion, and access to open reporting channels suggests that leadership genuinely values employee wellbeing as part of the ethical framework rather than simply as a liability management strategy.")

add_bold_label(doc, "b. Concerns: ", "One potential concern is that the social media policy, while understandable from a business perspective, could be interpreted broadly enough to discourage employees from speaking publicly about working conditions or other protected activities. The confidentiality policy is similarly broad and may create ambiguity about what partners can and cannot discuss regarding their own workplace experiences. These ambiguities could limit transparency or suppress legitimate employee speech.")

add_bold_label(doc, "c. Would You Work Here: ", "Based on the code of ethics alone, Starbucks appears to be a values-driven organization that takes employee dignity and ethical conduct seriously. The inclusive language, the emphasis on doing the right thing, and the accessibility of reporting mechanisms all suggest a workplace that aspires to genuine ethical standards. I would be open to working at Starbucks based on what this document communicates about organizational culture and values.")

add_bold_label(doc, "d. Additional Information: ", "Before making a final decision, I would want to understand how consistently the code is actually enforced, particularly in situations involving managers or senior leaders accused of misconduct. Concrete data on the outcomes of past ethics complaints and the strength of whistleblower protections in practice would be essential to assess. Independent employee feedback and reviews from platforms like Glassdoor would help determine whether the values expressed in the code reflect the day-to-day reality of working at the company.")

path = r"C:\Users\Jaylen.Davis\OneDrive - Southwestern College\Desktop\DoWhatever\projects\school\sophia-learning\business-ethics\touchstone-2"
os.makedirs(path, exist_ok=True)
filepath = os.path.join(path, "response.docx")
doc.save(filepath)
print("Saved:", filepath, "| Size:", os.path.getsize(filepath))
